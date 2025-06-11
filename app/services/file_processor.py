import os
import io
import uuid
import logging
from typing import Dict, List, Optional, Tuple, Any
from pathlib import Path
import mimetypes

# Document processing libraries
try:
    import PyPDF2
    import pdfplumber
except ImportError:
    PyPDF2 = None
    pdfplumber = None

try:
    import docx
except ImportError:
    docx = None

try:
    import openpyxl
    import pandas as pd
except ImportError:
    openpyxl = None
    pd = None

logger = logging.getLogger(__name__)


class FileProcessor:
    """Service for processing uploaded files and extracting text content"""
    
    SUPPORTED_FORMATS = {
        'text/plain': ['txt', 'md', 'markdown'],
        'text/html': ['html', 'htm'],
        'text/csv': ['csv'],
        'application/pdf': ['pdf'],
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': ['docx'],
        'application/vnd.ms-excel': ['xls'],
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': ['xlsx'],
        'application/json': ['json'],
        'application/xml': ['xml'],
        'text/xml': ['xml']
    }
    
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
    
    def __init__(self, upload_dir: str = "uploads"):
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(exist_ok=True)
    
    def is_supported_file(self, filename: str, content_type: str = None) -> bool:
        """Check if file format is supported"""
        extension = Path(filename).suffix.lower().lstrip('.')
        
        # Check by extension
        for supported_extensions in self.SUPPORTED_FORMATS.values():
            if extension in supported_extensions:
                return True
        
        # Check by content type
        if content_type and content_type in self.SUPPORTED_FORMATS:
            return True
        
        return False
    
    def validate_file(self, content: bytes, filename: str, content_type: str = None) -> Dict[str, Any]:
        """Validate uploaded file"""
        errors = []
        
        # Check file size
        if len(content) > self.MAX_FILE_SIZE:
            errors.append(f"File size ({len(content)} bytes) exceeds maximum allowed size ({self.MAX_FILE_SIZE} bytes)")
        
        # Check if file is supported
        if not self.is_supported_file(filename, content_type):
            errors.append(f"File format not supported. Supported formats: {', '.join(sum(self.SUPPORTED_FORMATS.values(), []))}")
        
        # Check if file is empty
        if len(content) == 0:
            errors.append("File is empty")
        
        return {
            "valid": len(errors) == 0,
            "errors": errors,
            "size": len(content),
            "filename": filename,
            "content_type": content_type
        }
    
    async def save_file(self, content: bytes, filename: str, tenant_id: str) -> str:
        """Save uploaded file to disk"""
        # Create tenant-specific directory
        tenant_dir = self.upload_dir / tenant_id
        tenant_dir.mkdir(exist_ok=True)
        
        # Generate unique filename
        file_id = str(uuid.uuid4())
        extension = Path(filename).suffix
        saved_filename = f"{file_id}{extension}"
        file_path = tenant_dir / saved_filename
        
        # Save file
        with open(file_path, 'wb') as f:
            f.write(content)
        
        logger.info(f"File saved: {filename} -> {file_path}")
        return str(file_path)
    
    async def extract_text_from_pdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF file"""
        if not PyPDF2 and not pdfplumber:
            raise ImportError("PDF processing requires PyPDF2 or pdfplumber: pip install PyPDF2 pdfplumber")
        
        text_content = ""
        metadata = {"pages": 0, "method": "unknown"}
        
        try:
            # Try pdfplumber first (better text extraction)
            if pdfplumber:
                with pdfplumber.open(file_path) as pdf:
                    metadata["pages"] = len(pdf.pages)
                    metadata["method"] = "pdfplumber"
                    
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += page_text + "\n\n"
            
            # Fallback to PyPDF2
            elif PyPDF2:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata["pages"] = len(pdf_reader.pages)
                    metadata["method"] = "PyPDF2"
                    
                    for page in pdf_reader.pages:
                        text_content += page.extract_text() + "\n\n"
        
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {e}")
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
        
        return text_content.strip(), metadata
    
    async def extract_text_from_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX file"""
        if not docx:
            raise ImportError("DOCX processing requires python-docx: pip install python-docx")
        
        try:
            doc = docx.Document(file_path)
            
            text_content = ""
            metadata = {"paragraphs": 0, "tables": 0}
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n\n"
                    metadata["paragraphs"] += 1
            
            # Extract tables
            for table in doc.tables:
                metadata["tables"] += 1
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content += " | ".join(row_text) + "\n"
                text_content += "\n"
        
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
        
        return text_content.strip(), metadata
    
    async def extract_text_from_excel(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from Excel file"""
        if not pd or not openpyxl:
            raise ImportError("Excel processing requires pandas and openpyxl: pip install pandas openpyxl")
        
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            text_content = ""
            metadata = {"sheets": len(excel_file.sheet_names), "rows": 0}
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                text_content += f"Sheet: {sheet_name}\n"
                text_content += "=" * 50 + "\n"
                
                # Convert DataFrame to text
                text_content += df.to_string(index=False) + "\n\n"
                metadata["rows"] += len(df)
        
        except Exception as e:
            logger.error(f"Error extracting text from Excel {file_path}: {e}")
            raise ValueError(f"Failed to extract text from Excel: {str(e)}")
        
        return text_content.strip(), metadata
    
    async def extract_text_from_csv(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from CSV file"""
        try:
            if pd:
                df = pd.read_csv(file_path)
                text_content = df.to_string(index=False)
                metadata = {"rows": len(df), "columns": len(df.columns)}
            else:
                # Fallback to basic reading
                with open(file_path, 'r', encoding='utf-8') as f:
                    text_content = f.read()
                metadata = {"rows": len(text_content.split('\n')), "method": "basic"}
        
        except Exception as e:
            logger.error(f"Error extracting text from CSV {file_path}: {e}")
            raise ValueError(f"Failed to extract text from CSV: {str(e)}")
        
        return text_content.strip(), metadata
    
    async def extract_text_from_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from plain text files"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'iso-8859-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text_content = f.read()
                    
                    metadata = {
                        "encoding": encoding,
                        "lines": len(text_content.split('\n')),
                        "characters": len(text_content)
                    }
                    
                    return text_content.strip(), metadata
                
                except UnicodeDecodeError:
                    continue
            
            raise ValueError("Could not decode file with any supported encoding")
        
        except Exception as e:
            logger.error(f"Error extracting text from text file {file_path}: {e}")
            raise ValueError(f"Failed to extract text from file: {str(e)}")
    
    async def process_file(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Process uploaded file and extract text content"""
        extension = Path(filename).suffix.lower().lstrip('.')
        
        try:
            if extension == 'pdf':
                text, metadata = await self.extract_text_from_pdf(file_path)
            elif extension == 'docx':
                text, metadata = await self.extract_text_from_docx(file_path)
            elif extension in ['xls', 'xlsx']:
                text, metadata = await self.extract_text_from_excel(file_path)
            elif extension == 'csv':
                text, metadata = await self.extract_text_from_csv(file_path)
            elif extension in ['txt', 'md', 'markdown', 'html', 'htm', 'json', 'xml']:
                text, metadata = await self.extract_text_from_text(file_path)
            else:
                raise ValueError(f"Unsupported file format: {extension}")
            
            # Basic text processing
            text = self.clean_text(text)
            
            return {
                "success": True,
                "text": text,
                "metadata": {
                    **metadata,
                    "filename": filename,
                    "extension": extension,
                    "file_path": file_path,
                    "text_length": len(text),
                    "word_count": len(text.split()) if text else 0
                }
            }
        
        except Exception as e:
            logger.error(f"Error processing file {filename}: {e}")
            return {
                "success": False,
                "error": str(e),
                "filename": filename,
                "file_path": file_path
            }
    
    def clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if line:  # Skip empty lines
                cleaned_lines.append(line)
        
        # Join with single newlines and normalize spaces
        cleaned_text = '\n'.join(cleaned_lines)
        
        # Remove excessive spaces
        import re
        cleaned_text = re.sub(r' +', ' ', cleaned_text)
        
        return cleaned_text.strip()
    
    async def delete_file(self, file_path: str) -> bool:
        """Delete uploaded file"""
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.info(f"File deleted: {file_path}")
                return True
            return False
        except Exception as e:
            logger.error(f"Error deleting file {file_path}: {e}")
            return False
    
    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """Get information about a file"""
        if not os.path.exists(file_path):
            return {"exists": False}
        
        stat = os.stat(file_path)
        return {
            "exists": True,
            "size": stat.st_size,
            "created": stat.st_ctime,
            "modified": stat.st_mtime,
            "path": file_path
        }


# Global instance
file_processor = FileProcessor() 